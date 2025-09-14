"""
CV service for handling CV queries using RAG (Retrieval-Augmented Generation).
"""
import uuid
from typing import List, Dict
from sqlalchemy.orm import Session
from fastapi import UploadFile
from PyPDF2 import PdfReader
from docx import Document

from app.models.cv import CVDocument
from app.core.vectorstore import add_documents, query_similar
from app.core.ai_client import OpenRouterClient

def _extract_text(file: UploadFile) -> str:
    """Extract text from uploaded files (PDF, DOCX, or text)."""
    if file.filename.lower().endswith(".pdf"):
        reader = PdfReader(file.file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif file.filename.lower().endswith(".docx"):
        d = Document(file.file)
        return "\n".join([p.text for p in d.paragraphs])
    else:
        return file.file.read().decode("utf-8", errors="ignore")

async def process_cv_upload(db: Session, file: UploadFile):
    """Process and store uploaded CV documents."""
    content = _extract_text(file)
    doc_id = str(uuid.uuid4())

    # Store meta in DB (optional)
    fpath = f"uploads/{doc_id}_{file.filename}"
    with open(fpath, "wb") as out:
        file.file.seek(0)
        out.write(file.file.read())

    cv = CVDocument(
        id=doc_id,
        filename=file.filename,
        file_path=fpath,
        file_size=len(content.encode("utf-8")),
        processed=True,
    )
    db.add(cv)
    db.commit()

    # Add to vector store along with your portfolio notes
    chunks = [content[i:i+1000] for i in range(0, len(content), 1000)]
    docs = [(f"{doc_id}_{i}", chunk) for i, chunk in enumerate(chunks)]
    await add_documents(docs)
    return {"message": "CV processed successfully"}

async def query_cv(question: str, detailed: bool = False) -> Dict:
    contexts = await query_similar(question, k=4)
    context_blob = "\n\n".join(contexts)

    system = (
        "You are Daniyal Ahmad's CV assistant. Answer using ONLY the provided context. "
        "If the answer is not present, reply: 'I don't have that information about Daniyal.'"
    )

    user_prompt = f"Context:\n{context_blob}\n\nQuestion: {question}\n"
    if detailed:
        user_prompt += "Provide a thorough answer, include examples and specifics if available."
        max_tokens = 800
    else:
        user_prompt += "Answer concisely in 1-3 sentences. If necessary, expand on request."
        max_tokens = 220

    messages = [{"role": "system", "content": system}, {"role": "user", "content": user_prompt}]
    ai_client = OpenRouterClient()
    answer = await ai_client.get_chat_response(messages, is_cv_query=True)
    confidence = 0.6 + 0.1 * min(len(contexts), 4)
    return {"answer": answer, "confidence": min(confidence, 0.95), "sources": contexts}

def _get_fallback_cv_response(question: str) -> str:
    """Generate a fallback response when AI is unavailable."""
    question_lower = question.lower()
    
    if "github" in question_lower or "username" in question_lower:
        return "Daniyal's GitHub username is **daniyalareeb** and you can find his projects at https://github.com/daniyalareeb."
    
    elif "project" in question_lower:
        return "Daniyal has built several impressive projects including this portfolio website, an AI Doctor system, a Crypto Gateway, and a Proxmox VE home lab."
    
    elif "skill" in question_lower or "experience" in question_lower:
        return "Daniyal is a skilled backend developer with expertise in Python, FastAPI, LLMs, and RAG systems. He currently works as a Backend Developer at Manage Your Sales."
    
    else:
        return "Daniyal Ahmad is a dedicated backend engineer with strong expertise in FastAPI, AI/ML technologies, and modern web development."

def refresh_cv():
    """
    Re-index CV data in Chroma vector store.
    This function can be called to refresh/rebuild the CV index.
    """
    try:
        # Import and run the setup script to refresh CV data
        import subprocess
        import sys
        from pathlib import Path
        
        setup_script = Path(__file__).parent.parent.parent / "scripts" / "setup_db.py"
        if setup_script.exists():
            result = subprocess.run([sys.executable, str(setup_script)], 
                                  capture_output=True, text=True, cwd=str(setup_script.parent.parent))
            if result.returncode == 0:
                print("✅ CV data refreshed successfully!")
                return True
            else:
                print(f"❌ CV refresh failed: {result.stderr}")
                return False
        else:
            print("❌ Setup script not found")
            return False
            
    except Exception as e:
        print(f"Error refreshing CV: {e}")
        raise e
